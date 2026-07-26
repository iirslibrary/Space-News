# -*- coding: utf-8 -*-
"""
IIRS Space Digest - Combined Production Version
Generates:
1) HTML digest
2) DOCX digest
Daily automated space news for IIRS employees
LAST 24 HOURS ROLLING WINDOW
"""

# =========================
# Imports
# =========================
import os
import re
import html
import time
import email.utils
import requests
import feedparser
import json
import hashlib

from io import BytesIO
from datetime import datetime, date, timedelta, timezone
from urllib.parse import urlparse, parse_qs, unquote, urljoin
from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode

from bs4 import BeautifulSoup
from newspaper import Article, Config
from googlenewsdecoder import gnewsdecoder

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


print("🚀 Starting Space News - LAST 24 HOURS WINDOW...")

# =========================
# Filters and Feed Lists
# =========================

EXCLUDED_KEYWORDS = r'(?i)(rape|murder|KYC|digilocker|arrest|crime|FIR|strikes|Rajya Sabha|Muslims|metro)'

REGIONAL_KEYWORDS = r'(?i)(space|satellite|remote sensing|gis|iirs|rrsc|nrsc|earth observation|glacier|landslide|cloudburst|disaster|floods|avalanche|earthquake|seismic|hyperspectral|air quality index| AQI |snowfall)'

NATIONAL_KEYWORDS = r'(?i)(isro|nrsc|nsil|chandrayaan| IIST |gaganyaan|pslv|glsv|lvm3|spadex|gsat|insat|resourcesat|cartosat|risat|launch|rocket|spacecraft|astronaut|shukrayaan|aditya|spaceport|sriharikota|indian space|vyommitra|eos|pslv-c62|axiom|nesac|nsss|sslv|nvs|hlvm3|om1)'

INTERNATIONAL_KEYWORDS = r'(?i)(nasa|esa|jaxa|cnsa|roscosmos|spacex|blue origin|artemis|starship|crew dragon|iss|international space station|hubble|james webb|mars rover|perseverance|insight|booster|orbital|launch|spacecraft|astronaut|spacewalk|satellite|mission|space agency)'

REGIONAL_FEEDS = [
    'https://www.amarujala.com/rss/uttarakhand.rss',
    'https://khabardevbhoomi.com/feed/',
    'https://devbhoomimedia.com/feed',
    'https://pioneeredge.in/feed',
    'https://www.livehindustan.com/uttarakhand/rss',
    'https://timesofindia.indiatimes.com/city/delhi/rssfeeds/1311474.cms',
    'https://indianexpress.com/section/cities/delhi/feed/',
    'https://www.hindustantimes.com/cities/delhi-news/rssfeed/',
]

yesterday = (date.today() - timedelta(days=1)).strftime('%Y-%m-%d')
google_isro = f'https://news.google.com/rss/search?q=ISRO+OR+NRSC+OR+IIRS+after:{yesterday}&hl=en-IN&gl=IN&-site:indianexpress.com&-site:thehindu.com&-site:timesofindia.indiatimes.com&-site:isro.gov.in&-site:economictimes.indiatimes.com'

NATIONAL_FEEDS = [
    'https://timesofindia.indiatimes.com/rssfeeds/1201659.cms',
    'https://indianexpress.com/section/science/feed/',
    'https://www.thehindu.com/sci-tech/science/rssfeed/',
    'https://www.thehindu.com/news/national/rssfeed/',
    'https://www.isro.gov.in/rssnews.xml',
    'https://government.economictimes.indiatimes.com/rss/digital-india',
    'https://government.economictimes.indiatimes.com/rss/policy',
    'https://government.economictimes.indiatimes.com/rss/governance',
    'https://government.economictimes.indiatimes.com/rss/smart-infra',
    'https://government.economictimes.indiatimes.com/rss/Defence',
    'https://government.economictimes.indiatimes.com/rss/economy',
    google_isro
]

INTERNATIONAL_FEEDS = [
    'https://www.esa.int/rss/rss-topnews.xml',
    'https://www.esa.int/rss/programmes.xml',
    'https://www.esa.int/rss/space_science.xml',
    'https://www.esa.int/rss/earth_observation.xml',
    'https://www.nasa.gov/rss/dyn/breaking_news.rss',
    'https://www.nasa.gov/rss/dyn/images_of_the_day.rss',
    'https://www.space.com/feeds/all',
    'https://spaceflightnow.com/feed/',
    'https://phys.org/rss-feed/space-news/',
    'https://www.thespacereview.com/rss.xml',
    'https://interestingengineering.com/feed',
]


# =========================
# Image Extraction Helpers
# =========================

BAD_IMAGE_HINTS = [
    "logo", "icon", "favicon", "sprite", "banner", "ads", "advert",
    "google-news", "gnews", "default", "placeholder", "avatar", 
    "author", "user", "profile", "category", "blank", "blank-image",
    "feedburner", "newsletter", "branding", "youtube", "facebook",
    "twitter", "instagram", "linkedin", "whatsapp", "telegram",
    "share", "social", "theme-assets", "thumb", "thumbnail", "small"
]

BAD_IMAGE_EXTENSIONS = [".svg", ".ico"]

def is_valid_image_url(url):
    if not url or not url.startswith("http"):
        return False
    low = url.lower()
    if any(low.endswith(ext) for ext in BAD_IMAGE_EXTENSIONS):
        return False
    if any(hint in low for hint in BAD_IMAGE_HINTS):
        return False
    if "/wp-content/themes/" in low:
        return False
    return True

def normalize_img_url(img, base_url):
    if not img:
        return None
    img = img.strip().replace("\\/", "/")
    if img.startswith("//"):
        img = "https:" + img
    elif img.startswith("/"):
        img = urljoin(base_url, img)
    return img

def score_image_url(img_url):
    if not img_url:
        return -999
    score = 0
    low = img_url.lower()
    if any(x in low for x in ["og:image", "og-image"]):
        score += 30
    if any(x in low for x in ["hero", "featured", "lead", "main", "article"]):
        score += 20
    if any(x in low for x in ["thumb", "thumbnail", "small", "icon", "logo", "sprite"]):
        score -= 40
    if any(x in low for x in ["120x", "150x", "180x", "200x", "300x"]):
        score -= 25
    if any(ext in low for ext in [".jpg", ".jpeg", ".png", ".webp"]):
        score += 5
    return score

def resolve_google_news_url(url):
    if not url or "news.google.com" not in url:
        return url
    try:
        decoded = gnewsdecoder(url)
        if isinstance(decoded, dict) and decoded.get("status"):
            decoded_url = decoded.get("decoded_url")
            if decoded_url and decoded_url.startswith("http"):
                return decoded_url
    except:
        pass
    return url

def resolve_msn_original_url(url):
    if not url or 'msn.com' not in url:
        return url
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0',
            'Accept-Language': 'en-US,en;q=0.9'
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        soup = BeautifulSoup(html_text, 'html.parser')

        canonical = soup.find('link', rel='canonical')
        if canonical and canonical.get('href'):
            canon_url = canonical['href'].strip()
            if canon_url.startswith('http') and 'msn.com' not in canon_url:
                return canon_url

        og_url = soup.find('meta', attrs={'property': 'og:url'})
        if og_url and og_url.get('content'):
            og_val = og_url['content'].strip()
            if og_val.startswith('http') and 'msn.com' not in og_val:
                return og_val

        for tag in soup.find_all('meta'):
            for attr in ['content', 'value']:
                val = tag.get(attr)
                if val and isinstance(val, str) and val.startswith('http'):
                    if 'msn.com' not in val and 'assets.msn.com' not in val and 'static.msn.com' not in val:
                        return val.strip()

        candidates = re.findall(r'https?://[^\s"\'<>\\]+', html_text)
        for cand in candidates:
            cand = unquote(cand.strip())
            if (cand.startswith('http') and 'msn.com' not in cand and 'assets.msn.com' not in cand
                and 'static.msn.com' not in cand and not any(x in cand.lower() for x in ['facebook.com', 'twitter.com', 'instagram.com', 'youtube.com'])):
                return cand

        parsed = urlparse(url)
        query = parse_qs(parsed.query)
        for key in ['url', 'src', 'source', 'redirect', 'u']:
            if key in query:
                possible = unquote(query[key][0])
                if possible.startswith('http') and 'msn.com' not in possible:
                    return possible

    except Exception:
        pass
    return url

def resolve_final_article_url(url):
    if not url:
        return url
    if 'news.google.com' in url:
        url = resolve_google_news_url(url)
    if 'msn.com' in url:
        url = resolve_msn_original_url(url)
    return url

def extract_image_from_raw_html(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0",
            "Accept-Language": "en-US,en;q=0.9"
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        soup = BeautifulSoup(html_text, 'html.parser')
        
        meta_tags = soup.find_all('meta')
        target_properties = ['og:image', 'og:image:url']
        target_names = ['twitter:image', 'twitter:image:src']

        for tag in meta_tags:
            prop = tag.get('property', '').lower()
            name = tag.get('name', '').lower()
            
            if prop in target_properties or name in target_names:
                img = tag.get('content')
                if img:
                    img = img.strip()
                    if img.startswith("//"):
                        img = "https:" + img
                    elif img.startswith("/"):
                        img = urljoin(url, img)
                        
                    if is_valid_image_url(img):
                        return img

        patterns = [
            r'<img[^>]+data-lazy-src=["\']([^"\']+)["\']',
            r'<img[^>]+data-src=["\']([^"\']+)["\']',
            r'<img[^>]+data-srcset=["\']([^"\']+)["\']',
            r'<img[^>]+src=["\']([^"\']+)["\']'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, html_text, flags=re.I)
            for match in matches:
                img = match.strip().split()[0] 
                
                if img.startswith("//"):
                    img = "https:" + img
                elif img.startswith("/"):
                    img = urljoin(url, img)
                    
                if is_valid_image_url(img):
                    return img
    except Exception:
        pass
    return None

def extract_image_from_jsonld_or_scripts(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0",
            "Accept-Language": "en-US,en;q=0.9"
        }
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        html_text = response.text

        patterns = [
            r'"image"\s*:\s*"([^"]+)"',
            r'"thumbnailUrl"\s*:\s*"([^"]+)"',
            r'"contentUrl"\s*:\s*"([^"]+)"',
            r'"url"\s*:\s*"([^"]+\.(?:jpg|jpeg|png|webp))"'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, html_text, flags=re.I)
            for img in matches:
                img = img.replace("\\/", "/")
                if img.startswith("//"):
                    img = "https:" + img
                elif img.startswith("/"):
                    img = urljoin(url, img)
                if is_valid_image_url(img):
                    return img
    except:
        pass
    return None

def extract_image_with_newspaper(url):
    if not url or not url.startswith("http"):
        return None
    try:
        config = Config()
        config.browser_user_agent = 'Mozilla/5.0'
        config.request_timeout = 20
        article = Article(url, config=config)
        article.download()
        article.parse()

        if article.top_image and article.top_image.startswith("http") and is_valid_image_url(article.top_image):
            return article.top_image

        if article.images:
            images = []
            for img in article.images:
                if isinstance(img, str) and img.startswith("http") and is_valid_image_url(img):
                    images.append(img)
            if images:
                images = sorted(images, key=score_image_url, reverse=True)
                return images[0]
    except:
        pass
    return None

def extract_first_image_url(entry, article_url=None):
    # 1. PRIORITY 1: Always check the clean RSS feed data first
    try:
        for item in entry.get("media_content", []):
            url = item.get("url")
            if is_valid_image_url(url): return url
    except Exception: pass

    try:
        for item in entry.get("media_thumbnail", []):
            url = item.get("url")
            if is_valid_image_url(url): return url
    except Exception: pass

    try:
        for link in entry.get("links", []):
            href = link.get("href", "")
            link_type = link.get("type", "")
            rel = link.get("rel", "")
            if href and href.startswith("http") and (rel == "enclosure" or str(link_type).startswith("image/")):
                if is_valid_image_url(href): return href
    except Exception: pass

    # 2. PRIORITY 2: Fallback to web scraping only if the RSS feed has no image attached
    article_url = resolve_final_article_url(article_url) if article_url else None
    if article_url:
        image = extract_image_from_raw_html(article_url)
        if image: return image
        image = extract_image_from_jsonld_or_scripts(article_url)
        if image: return image
        image = extract_image_with_newspaper(article_url)
        if image: return image

    return None

def try_download_image(image_url, timeout=20):
    if not image_url: return None
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(image_url, headers=headers, timeout=timeout)
        response.raise_for_status()
        content_type = response.headers.get("Content-Type", "").lower()
        if "image" not in content_type: return None
        return BytesIO(response.content)
    except Exception: return None

# =========================
# Text Helpers
# =========================
def clean_source_name(source_name):
    if not source_name: return "Unknown Source"
    source_name = html.unescape(str(source_name))
    return re.sub(r'\.\.\.$', '', source_name).strip()

def sanitize_filename(name):
    return re.sub(r'[^a-zA-Z0-9_-]+', '_', name)

def normalize_text(text):
    if not text: return ""
    text = html.unescape(str(text))
    replacements = { "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"', "\u2013": "-", "\u2014": "-", "\u00a0": " ", "\u200b": "", "\ufeff": "", "\\|": "|", "\\'": "'", '\\"': '"' }
    for old, new in replacements.items(): text = text.replace(old, new)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def clean_body_text(text, title=""):
    text = normalize_text(text)
    if not text: return ""
    lines = [ln.strip() for ln in text.splitlines()]
    bad_phrases = ["your browser does not support javascript", "related articles", "add asianet newsable as a preferred source", "google news", "follow us on", "read more", "advertisement", "recommended stories", "suggested articles", "share this article", "click here"]
    cleaned = []
    seen = set()
    for ln in lines:
        if not ln: continue
        low = ln.lower().strip()
        if any(bp in low for bp in bad_phrases): continue
        if title and low == normalize_text(title).lower(): continue
        if len(low) < 3: continue
        if low in seen: continue
        seen.add(low)
        cleaned.append(ln)
    text = "\n".join(cleaned)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text).strip()
    return text

def split_into_paragraphs(text):
    text = clean_body_text(text)
    if not text: return []
    paras = re.split(r'\n{2,}', text)
    final_paras = []
    for para in paras:
        para = re.sub(r'\s+', ' ', para).strip()
        if not para: continue
        if len(para) < 20: continue
        final_paras.append(para)
    return final_paras

def sanitize_html_content(text):
    if not text: return ''
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text[:380] + '...' if len(text) > 380 else text.strip()


# =========================
# News Timing
# =========================
def is_within_last_24_hours(entry):
    now = datetime.now(timezone.utc)
    cutoff_time = now - timedelta(hours=24)
    pub_struct = entry.get('published_parsed') or entry.get('updated_parsed') or entry.get('created_parsed')
    if pub_struct:
        try:
            pub_time = datetime(*pub_struct[:6], tzinfo=timezone.utc)
            return pub_time >= cutoff_time
        except: pass
    date_str = entry.get('published') or entry.get('updated') or entry.get('created')
    if date_str:
        try:
            parsed_tuple = email.utils.parsedate_tz(date_str)
            if parsed_tuple:
                ts = email.utils.mktime_tz(parsed_tuple)
                pub_time = datetime.fromtimestamp(ts, timezone.utc)
                return pub_time >= cutoff_time
        except: pass
    return False


# =========================
# Feed Fetching
# =========================
def fetch_news_from_feeds(feeds, max_articles=6):
    news = []
    seen_image_urls = set()

    for url in feeds:
        try:
            feed = feedparser.parse(url)
            print(f"📱 {feed.feed.get('title', 'Unknown')} - checking...")
            for entry in feed.entries[:15]:
                if not is_within_last_24_hours(entry): continue
                title_lower = entry.title.lower()
                raw_summary = entry.get('summary', '') or entry.get('description', '')
                summary_lower = raw_summary.lower()
                full_text_check = title_lower + " " + summary_lower

                if url in REGIONAL_FEEDS: keyword_pattern = REGIONAL_KEYWORDS
                elif url in NATIONAL_FEEDS: keyword_pattern = NATIONAL_KEYWORDS
                else: keyword_pattern = INTERNATIONAL_KEYWORDS

                if not re.search(keyword_pattern, title_lower): continue
                if re.search(EXCLUDED_KEYWORDS, full_text_check):
                    print(f"🗑️ REMOVED (Excluded content): {entry.title[:40]}...")
                    continue

                original_link = entry.link
                final_link = resolve_final_article_url(original_link)
                image_url = extract_first_image_url(entry, final_link)
                
                if image_url in seen_image_urls:
                    print(f"🚫 BLOCKED: Image already used by a previous article (Sidebar Hijack). Setting to null.")
                    image_url = None
                if image_url: seen_image_urls.add(image_url)

                summary = sanitize_html_content(raw_summary)
                title = re.sub(r'<[^>]+>', '', entry.title)
                news.append({ 'title': title, 'link': final_link, 'source': feed.feed.get('title', ''), 'summary': summary, 'image': image_url, 'ai_label': None })

                print(f"✅ NEW (24h): {title[:60]}...")
                print(f"🔗 Final link: {final_link}")
                print(f"🖼️ Image assigned: {image_url}")
                if len(news) >= max_articles: break
            if len(news) >= max_articles: break
        except Exception as e: print(f"⚠️ Skip {url}: {e}")
    return news

def get_ist_today():
    ist = timezone(timedelta(hours=5, minutes=30))
    return datetime.now(ist).strftime("%Y-%m-%d")

published_state = {}
try:
    with open("published_digest_state.json", "r", encoding="utf-8") as f:
        published_state = json.load(f)
except FileNotFoundError:
    published_state = {}

published_for_date = str(published_state.get("published_for_date", "")).strip()
is_finalized = (published_for_date == get_ist_today())

def make_articles_html(news_list, is_finalized=False):
    html_out = ""
    for i, item in enumerate(news_list, 1):
        article_url = resolve_final_article_url(normalize_text(item.get("link", "")))
        safe_url = html.escape(article_url, quote=True)
        image_html = ''
        if item.get("image"):
            image_html = f'<img src="{html.escape(item["image"], quote=True)}" alt=" image" class="card-image" loading="lazy" onerror="this.style.display=\'none\'">'
        
        flag_html = ""
        if not is_finalized:
            flag_html = f'<label class="flag-item"><input type="checkbox" class="flag-checkbox" value="{safe_url}">Flag this article</label>'
        
        raw_ai_label = normalize_text(item.get("ai_label", "")).strip()
        ai_label_lower = raw_ai_label.lower()
        if "high" in ai_label_lower: ai_class, ai_label = "ai-high", "Highly Relevant"
        elif "medium" in ai_label_lower: ai_class, ai_label = "ai-medium", "Medium Relevant"
        else: ai_class, ai_label = "ai-low", "Not Relevant"

        ai_html = f'<div class="ai-relevance-box hidden-ai"><span class="ai-relevance-pill {ai_class}">{html.escape(ai_label)}</span></div>'

        html_out += f'''
            <div class="news-card">
                <div class="card-content">
                    {image_html}
                    <div class="card-title">
                        <a href="{safe_url}" target="_blank" rel="noopener noreferrer">{i}. {item["title"]}</a>
                    </div>
                    <div class="card-summary">{item["summary"]}</div>
                    <div class="card-actions">
                        <div class="card-action-left">
                            <a class="read-more" href="{safe_url}" target="_blank" rel="noopener noreferrer">Read Full Article →</a>
                        </div>
                        <div class="card-action-center">{ai_html}</div>
                        <div class="card-action-right">{flag_html}</div>
                    </div>
                </div>
            </div>
        '''

    if not is_finalized:
        html_out += '''
            <div class="bottom-actions">
                <button type="button" class="flag-submit-btn" onclick="submitFlags()">Submit Flagged Articles</button>
                <button type="button" class="publish-btn" onclick="publishCurrentList()">Publish</button>
            </div>
        '''
    return html_out

# =========================
# DOCX Helpers
# =========================
def add_bottom_border(paragraph):
    p, pPr = paragraph._p, paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), 'A6A6A6')
    pbdr.append(bottom); pPr.append(pbdr)

def add_top_border(paragraph, color="D9D9D9", size="6", space="4"):
    p, pPr = paragraph._p, paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), size)
    top.set(qn('w:space'), space)
    top.set(qn('w:color'), color)
    pbdr.append(top); pPr.append(pbdr)

def add_box_border(paragraph, color="808080", size="8", space="8"):
    p, pPr = paragraph._p, paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for side_name in ['top', 'left', 'bottom', 'right']:
        side = OxmlElement(f'w:{side_name}')
        side.set(qn('w:val'), 'single')
        side.set(qn('w:sz'), size)
        side.set(qn('w:space'), space)
        side.set(qn('w:color'), color)
        pbdr.append(side)
    pPr.append(pbdr)

def add_hyperlink(paragraph, text, url, color="0000FF", underline=True):
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run, rPr = OxmlElement("w:r"), OxmlElement("w:rPr")
    if color: c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single" if underline else "none"); rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t"); text_elem.text = text; new_run.append(text_elem)
    hyperlink.append(new_run); paragraph._p.append(hyperlink)
    return hyperlink

def set_section_columns(section, num_cols=1, space=360):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols: cols = cols[0]
    else: cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols)); cols.set(qn('w:space'), str(space))

def add_footer_to_section(section):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    if paragraph.runs:
        for run in paragraph.runs: run.text = ""
    add_top_border(paragraph, color="D9D9D9", size="6", space="4")
    run1 = paragraph.add_run("पुस्तकालय एवं सूचना संसाधन प्रभाग द्वारा संकलित, भा.सु.सं.सं")
    run1.font.name, run1.font.size = "Mangal", Pt(9)
    run1.add_break()
    run2 = paragraph.add_run("Compiled by Library and Information Resource Division, IIRS")
    run2.font.name, run2.font.size = "Times New Roman", Pt(9)

def apply_footer_to_all_sections(doc):
    for section in doc.sections: add_footer_to_section(section)

def fetch_full_article_text(url, fallback_summary="", title=""):
    fallback_summary = clean_body_text(fallback_summary, title=title)
    if not url or url == '#': return fallback_summary
    url = resolve_final_article_url(url)
    try:
        config = Config(); config.browser_user_agent = 'Mozilla/5.0'; config.request_timeout = 20
        article = Article(url, config=config); article.download(); article.parse()
        text = clean_body_text(article.text or '', title=title)
        if len(text) >= 300: return text
    except: pass

    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=20)
        response.raise_for_status()
        raw_html = response.text
        raw_html = re.sub(r'<script.*?>.*?</script>', ' ', raw_html, flags=re.I | re.S)
        raw_html = re.sub(r'<style.*?>.*?</style>', ' ', raw_html, flags=re.I | re.S)
        patterns = [ r'<article[^>]*>(.*?)</article>', r'<main[^>]*>(.*?)</main>', r'<div[^>]+class=["\'][^"\']*(?:article|story|content|main-content|post-content|entry-content|td-post-content|news-detail|story-detail)[^"\']*["\'][^>]*>(.*?)</div>' ]
        extracted = ''
        for pattern in patterns:
            matches = re.findall(pattern, raw_html, flags=re.I | re.S)
            if matches:
                flat = []
                for m in matches[:2]:
                    if isinstance(m, tuple): flat.extend([x for x in m if x])
                    else: flat.append(m)
                extracted = ' '.join(flat); break
        if not extracted: extracted = raw_html
        extracted = re.sub(r'</p>|<br\s*/?>|</div>|</section>|</article>|</li>|</h[1-6]>', '\n', extracted, flags=re.I)
        extracted = re.sub(r'<li[^>]*>', '- ', extracted, flags=re.I)
        extracted = re.sub(r'<[^>]+>', ' ', extracted)
        extracted = clean_body_text(extracted, title=title)
        if len(extracted) >= 300: return extracted
    except: pass
    return fallback_summary

def add_article_body_single_column(doc, paragraphs):
    if not paragraphs: paragraphs = ['Summary not available.']
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after, p.paragraph_format.line_spacing, p.alignment = Pt(6), 1.15, WD_ALIGN_PARAGRAPH.JUSTIFY
        para = re.sub(r'\s+', ' ', para).strip()
        run = p.add_run(para)
        run.font.name, run.font.size = 'Times New Roman', Pt(12)

def generate_docx(news_items, output_path, digest_date_str):
    doc = Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.6), Inches(0.6)
    section.left_margin, section.right_margin = Inches(0.7), Inches(0.7)
    set_section_columns(section, num_cols=1)
    
    styles = doc.styles
    styles['Normal'].font.name, styles['Normal'].font.size = 'Times New Roman', Pt(11)
    add_footer_to_section(section)

    logo_table = doc.add_table(rows=1, cols=3); logo_table.autofit = False
    cells = logo_table.rows[0].cells
    cells[0].width, cells[1].width, cells[2].width = Inches(1.2), Inches(4.8), Inches(1.2)
    
    left_p = cells[0].paragraphs[0]; left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try: left_p.add_run().add_picture("./assets/iirs.png", width=Inches(0.55))
    except: pass
    
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    right_p = cells[2].paragraphs[0]; right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    try: right_p.add_run().add_picture("./assets/isro-logo-png_seeklogo-304812.png", width=Inches(0.55))
    except: pass

    doc.add_paragraph()
    header_box = doc.add_paragraph()
    header_box.alignment, header_box.paragraph_format.space_before, header_box.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(3), Pt(10)
    
    title_run = header_box.add_run("🌌 अंतरिक्ष समाचार | Space News")
    title_run.bold, title_run.font.name, title_run.font.size = True, "Times New Roman", Pt(13)
    title_run.add_break()
    
    date_run = header_box.add_run(digest_date_str)
    date_run.bold, date_run.font.name, date_run.font.size = False, "Times New Roman", Pt(10)
    add_box_border(header_box, color="808080", size="8", space="8")
    doc.add_paragraph()

    from PIL import Image
    import io

    for idx, item in enumerate(news_items, start=1):
        title = normalize_text(item.get('title', 'Untitled'))
        link = resolve_final_article_url(normalize_text(item.get('link', '')))
        summary = normalize_text(item.get('summary', ''))
        image_url = item.get('image')

        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{idx}. {title}')
        run.bold, run.font.name, run.font.size = True, 'Times New Roman', Pt(13)

        if link and link != '#':
            link_p = doc.add_paragraph(); link_p.paragraph_format.space_after, link_p.alignment = Pt(4), WD_ALIGN_PARAGRAPH.LEFT
            add_hyperlink(link_p, "Read more", link)

            image_stream = try_download_image(image_url)
            if image_stream:
                inserted = False
                try:
                    image_stream.seek(0)
                    img = Image.open(image_stream)
                    img.load()
                    if img.mode in ("RGBA", "P"): img = img.convert("RGB")
                    clean_stream = io.BytesIO(); img.save(clean_stream, format="PNG"); clean_stream.seek(0)
                    img_p = doc.add_paragraph(); img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_p.add_run().add_picture(clean_stream, width=Inches(4.8))
                    img_p.paragraph_format.space_after = Pt(6); inserted = True
                except Exception as e: print(f"Clean image conversion failed: {e}")

                if not inserted:
                    try:
                        image_stream.seek(0)
                        img_p = doc.add_paragraph(); img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_p.add_run().add_picture(image_stream, width=Inches(4.8))
                        img_p.paragraph_format.space_after = Pt(6)
                    except Exception as e: print(f"Original image insert failed: {e}")

        body_text = fetch_full_article_text(url=link, fallback_summary=summary, title=title)
        body_text = clean_body_text(body_text, title=title)
        body_paragraphs = split_into_paragraphs(body_text)

        if not body_paragraphs:
            fallback_clean = clean_body_text(summary, title=title)
            body_paragraphs = split_into_paragraphs(fallback_clean)

        add_article_body_single_column(doc, body_paragraphs)

        if idx != len(news_items):
            sep = doc.add_paragraph(); add_bottom_border(sep); doc.add_paragraph('')

    apply_footer_to_all_sections(doc)
    doc.save(output_path)
    print(f'DOCX saved: {output_path}')


FLAG_FILE = "flagged_urls.json"

def load_flagged_urls():
    if not os.path.exists(FLAG_FILE): return set()
    try:
        with open(FLAG_FILE, "r", encoding="utf-8") as f: data = json.load(f)
        if isinstance(data, list): return {url for url in data if url}
        if isinstance(data, dict): return {url for url in data.get("flagged_urls", []) if url}
        return set()
    except Exception as e:
        print(f"⚠️ Failed to load flagged URLs: {e}")
        return set()

def normalize_url_for_compare(url):
    if not url: return ""
    url = normalize_text(url).strip()
    if not url: return ""
    if not url.startswith(("http://", "https://")): url = "https://" + url
    try:
        parsed = urlsplit(url)
        scheme, netloc = "https", parsed.netloc.lower().strip()
        if netloc.startswith("www."): netloc = netloc[4:]
        if netloc.endswith(":80"): netloc = netloc[:-3]
        elif netloc.endswith(":443"): netloc = netloc[:-4]
        path = parsed.path or "/"
        while "//" in path: path = path.replace("//", "/")
        if path != "/" and path.endswith("/"): path = path[:-1]
        query_params = parse_qsl(parsed.query, keep_blank_values=False)
        filtered_params = [(k, v) for k, v in query_params if not k.lower().startswith("utm_") and k.lower() not in {"fbclid", "gclid", "mc_cid", "mc_eid", "igshid"}]
        return urlunsplit((scheme, netloc, path, urlencode(filtered_params, doseq=True), ""))
    except Exception: return url.strip()

def filter_flagged_news(news_items):
    flagged_urls = load_flagged_urls()
    if not flagged_urls: return news_items
    normalized_flagged = {normalize_url_for_compare(url) for url in flagged_urls if url}
    
    filtered_news = []
    for item in news_items:
        raw_link = normalize_text(item.get("link", ""))
        final_link = resolve_final_article_url(raw_link)
        if normalize_url_for_compare(final_link) not in normalized_flagged:
            item["link"] = final_link
            filtered_news.append(item)
        else: print(f"🚫 Removed flagged article: {final_link}")
    return filtered_news

# =========================
# Snapshot Loading & Deletion
# =========================
SNAPSHOT_DIR = Path("snapshots")
SNAPSHOT_DIR.mkdir(exist_ok=True)
TODAY = datetime.now().date()
TODAY_STR = TODAY.strftime("%Y-%m-%d")
TODAY_SNAPSHOT_FILE = SNAPSHOT_DIR / f"{TODAY_STR}.json"

def cleanup_old_snapshots(keep_days=45):
    cutoff_date = TODAY - timedelta(days=keep_days - 1)
    for file in SNAPSHOT_DIR.glob("*.json"):
        if file.name == "published_digest_state.json": continue
        try:
            if datetime.strptime(file.stem, "%Y-%m-%d").date() < cutoff_date: file.unlink()
        except ValueError: pass
    for file in Path(".").glob("Space_News_*.*"):
        if file.suffix not in ['.html', '.docx']: continue
        try:
            if datetime.strptime(file.stem.replace("Space_News_", ""), "%d_%m_%Y").date() < cutoff_date: file.unlink()
        except ValueError: pass

def load_snapshot_by_date(date_str):
    snapshot_file = SNAPSHOT_DIR / f"{date_str}.json"
    if snapshot_file.exists():
        try:
            with snapshot_file.open("r", encoding="utf-8") as f: return json.load(f)
        except Exception: pass
    return None

def save_snapshot_by_date(date_str, data):
    try:
        with (SNAPSHOT_DIR / f"{date_str}.json").open("w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception: pass

def snapshot_needs_ai_labels(news_items):
    return any(str(item.get("ai_label", "")).strip() not in {"Highly Relevant", "Medium Relevant", "Not Relevant"} for item in news_items)

def load_today_snapshot():
    if TODAY_SNAPSHOT_FILE.exists():
        try:
            with TODAY_SNAPSHOT_FILE.open("r", encoding="utf-8") as f: return json.load(f)
        except Exception: pass
    return None

def save_today_snapshot(all_news):
    try:
        with TODAY_SNAPSHOT_FILE.open("w", encoding="utf-8") as f: json.dump(all_news, f, indent=2, ensure_ascii=False)
    except Exception: pass

OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
VALID_AI_LABELS = {"Highly Relevant", "Medium Relevant", "Not Relevant"}

def make_article_key(item):
    return hashlib.sha1(" | ".join([normalize_text(item.get("title", "")) or "", normalize_text(item.get("source", "")) or "", normalize_text(item.get("link", "")) or ""]).encode("utf-8")).hexdigest()[:16]

def build_ai_payload(all_news):
    payload = []
    for item in all_news:
        if str(item.get("ai_label", "")).strip() in VALID_AI_LABELS: continue
        ai_key = make_article_key(item)
        item["_ai_key"] = ai_key
        payload.append({"key": ai_key, "title": normalize_text(item.get("title", "")) or "", "source": normalize_text(item.get("source", "")) or "", "summary": normalize_text(item.get("summary", "")) or "", "link": normalize_text(item.get("link", "")) or ""})
    return payload

def classify_articles_batch(payload):
    api_key = os.getenv("OPENROUTER_API_KEY")
    model_name = os.getenv("OPENROUTER_MODEL", "openrouter/free")
    if not payload: return []
    if not api_key: return [{"key": row["key"], "ai_label": "Medium Relevant"} for row in payload]

    schema = {
        "type": "object",
        "properties": { "results": { "type": "array", "items": { "type": "object", "properties": { "key": {"type": "string"}, "ai_label": { "type": "string", "enum": ["Highly Relevant", "Medium Relevant", "Not Relevant"] } }, "required": ["key", "ai_label"], "additionalProperties": False } } },
        "required": ["results"], "additionalProperties": False
    }
    
    system_prompt = "You are classifying news articles for an internal IIRS space-news digest.\nAssign exactly one label to each article: Highly Relevant, Medium Relevant, or Not Relevant."
    user_prompt = f"Classify the following article list.\n\n{json.dumps(payload, ensure_ascii=False, indent=2)}"

    try:
        response = requests.post(
            OPENROUTER_API_URL, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json", "HTTP-Referer": "https://iirslibrary.github.io", "X-Title": "IIRS Space News Digest"},
            json={"model": model_name, "temperature": 0, "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_schema", "json_schema": {"name": "article_relevance", "strict": True, "schema": schema}}},
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        parsed = json.loads(data["choices"][0]["message"]["content"])
        cleaned = []
        for row in parsed.get("results", []):
            if str(row.get("key", "")).strip() and str(row.get("ai_label", "")).strip() in VALID_AI_LABELS:
                cleaned.append({"key": str(row.get("key", "")).strip(), "ai_label": str(row.get("ai_label", "")).strip()})
        return cleaned
    except Exception: return [{"key": row["key"], "ai_label": "Medium Relevant"} for row in payload]

def merge_ai_labels(all_news, ai_results):
    label_map = {row["key"]: row["ai_label"] for row in ai_results if row.get("key") and row.get("ai_label") in VALID_AI_LABELS}
    for item in all_news:
        if str(item.get("ai_label", "")).strip() in VALID_AI_LABELS: item.pop("_ai_key", None); continue
        key = item.get("_ai_key")
        item["ai_label"] = label_map.get(key, "Medium Relevant")
        item.pop("_ai_key", None)
    return all_news

def enrich_articles_with_ai_relevance(all_news):
    if not all_news: return all_news
    payload = build_ai_payload(all_news)
    if not payload: return all_news
    ai_results = classify_articles_batch(payload)
    return merge_ai_labels(all_news, ai_results)

# =========================
# Main Fetch
# =========================
cleanup_old_snapshots(keep_days=60)
snapshot_override = os.getenv("SNAPSHOT_DATE", "").strip()

if snapshot_override:
    all_news = load_snapshot_by_date(snapshot_override)
    if all_news is None: raise FileNotFoundError(f"Requested snapshot not found: {snapshot_override}")
else: all_news = load_today_snapshot()

if all_news is None:
    regional_news = fetch_news_from_feeds(REGIONAL_FEEDS, max_articles=5)
    national_news = fetch_news_from_feeds(NATIONAL_FEEDS, max_articles=6)
    international_news = fetch_news_from_feeds(INTERNATIONAL_FEEDS, max_articles=8)

    all_news = []
    for news_list, category in [(regional_news, "🏔️ Regional Updates"), (national_news, "🇮🇳 National Updates"), (international_news, "🌌 International Updates")]:
        for item in news_list: item["category"] = category; all_news.append(item)

    if not all_news: all_news.append({"title": "No space news in last 24h", "link": "#", "source": "IIRS Digest", "summary": "Check back tomorrow!", "image": None, "category": "System", "ai_label": "Not Relevant"})
    all_news = enrich_articles_with_ai_relevance(all_news)
    save_today_snapshot(all_news)
else:
    if snapshot_needs_ai_labels(all_news):
        all_news = enrich_articles_with_ai_relevance(all_news)
        if snapshot_override: save_snapshot_by_date(snapshot_override, all_news)
        else: save_today_snapshot(all_news)

all_news = filter_flagged_news(all_news)

# # =========================
# # HTML Output (Staging vs Production setup)
# # =========================

# # Check if the workflow is running automatically or via manual publish trigger
# run_mode = os.environ.get('WORKFLOW_RUN_MODE', 'draft').lower()

# all_articles_html = make_articles_html(all_news, is_finalized=is_finalized)
# ist_offset = timezone(timedelta(hours=5, minutes=30))
# now_ist = datetime.now(ist_offset)
# digit_map = str.maketrans("0123456789", "०१२३४५६७८९")
# eng_date, eng_time, hindi_date = now_ist.strftime("%d-%m-%Y"), now_ist.strftime("%H:%M"), now_ist.strftime("%d-%m-%Y").translate(digit_map)
# timestamp = f"दिनांक: {hindi_date} • Date: {eng_date} • Time: {eng_time} IST • {len(all_news)} Updates"
# build_version = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

# # Your full interactive HTML Template (for public OR admin use depending on the mode)
# html_body = f"""<!DOCTYPE html>
# <html data-theme="dark">
# <head>
# <script src="https://accounts.google.com/gsi/client" async defer></script>
# <meta charset="UTF-8">
# <meta name="viewport" content="width=device-width, initial-scale=1.0">
# <meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">
# <meta http-equiv="Pragma" content="no-cache">
# <meta http-equiv="Expires" content="0">
# <meta name="build-version" content="{build_version}">
# <link rel="stylesheet" href="review-styles.css?v={build_version}">
# </head>
# <body>
# <button class="theme-toggle" id="themeToggle" title="Toggle Theme">☀️</button>
# <div class="scroll-container">
#     <div class="page-header">
#         <img src="./assets/iirs.png" alt="IIRS Logo" class="top-logo left-logo">
#         <div class="header-title-wrap">
            # <h2>🌌 अंतरिक्ष समाचार | Space News</h2>
            <h2>🌌 अंतरिक्ष समाचार <span class="title-separator">|</span> <span class="title-eng">Space News</span></h2>
#         </div>
#         <img src="./assets/ISRO-Color.svg" alt="ISRO Logo" class="top-logo right-logo">
#     </div>
#     <p style="text-align:center; color:var(--text-secondary); margin-top:12px; margin-bottom:40px;">
#         {timestamp}
#     </p>
#     <div class="auth-panel" id="reviewerAccessBox">
#         <div class="auth-panel-title">Reviewer Access</div>
#         <div class="google-btn-row">
#             <div id="googleSignInBtn"></div>
#         </div>
#         <div id="signedInUser" style="display:none;"></div>
#         <div id="authMessage" class="auth-message"></div>
#     </div>
#     <div class="ai-toggle-wrap" style="display:none;">
#         <button type="button" class="ai-toggle-btn" onclick="toggleAIRelevance()">
#             Show AI Relevance
#         </button>
#     </div>
#     {all_articles_html}
#     <footer class="footer">
#         <div class="footer-text">
#             <div>पुस्तकालय एवं सूचना संसाधन प्रभाग द्वारा संकलित, भा.सु.सं.सं</div>
#             <div>Compiled by Library and Information Resource Division, IIRS</div>
#         </div>
#     </footer>
# </div>
# <div id="customToast" class="custom-toast">
#     <div class="toast-title" id="toastTitle"></div>
#     <div class="toast-message" id="toastMessage"></div>
# </div>
# <script src="review-actions.js?v={build_version}"></script>
# </body>
# </html>
# """

# # The safe screen shown to regular employees during Draft Mode
# holding_screen_html = f"""<!DOCTYPE html>
# <html lang="en">
# <head>
#     <meta charset="UTF-8">
#     <meta name="viewport" content="width=device-width, initial-scale=1.0">
#     <title>Space News Digest - Coming Shortly</title>
#     <style>
#         body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; background-color: #f8fafc; color: #334155; text-align: center; }}
#         .card {{ background: white; padding: 2.5rem; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); max-width: 480px; border-top: 4px solid #3b82f6; }}
#         h1 {{ font-size: 1.5rem; margin-bottom: 0.5rem; color: #0f172a; }}
#         p {{ color: #64748b; font-size: 1rem; line-height: 1.6; margin-top: 15px; }}
#         .highlight-msg {{ font-weight: bold; color: #1e293b; margin-top: 20px; }}
#     </style>
# </head>
# <body>
#     <div class="card">
#         <h1>📡 Today's news is coming shortly!</h1>
#         <p>Please check back later!</p>
#         <p class="highlight-msg">Thanks for cooperation, WE WANT THIS MESSAGE TO DISPLAY</p>
#     </div>
# </body>
# </html>
# """

# html_filename = f'Space_News_{datetime.now().strftime("%Y%m%d")}.html'

# # File generation conditional based on the run mode
# if run_mode == 'publish':
#     # PUBLISH: Save real news to main file for the public
#     with open(html_filename, 'w', encoding='utf-8') as f:
#         f.write(html_body)
#     print(f"✅ [PUBLISH] SAVED official public news: {html_filename} with {len(all_news)} items")
# else:
#     # DRAFT: Hide real news from public, put it in admin-review.html
#     with open(html_filename, 'w', encoding='utf-8') as f:
#         f.write(holding_screen_html)
#     print(f"🔒 [DRAFT] SAVED safe holding screen for the public to {html_filename}")
    
#     with open("admin-review.html", "w", encoding="utf-8") as f:
#         f.write(html_body)
#     print(f"🛠️ [DRAFT] SAVED review tools and full news to admin-review.html")



# =========================
# HTML Output (Staging vs Production setup)
# =========================

# Check if the workflow is running automatically or via manual publish trigger
run_mode = os.environ.get('WORKFLOW_RUN_MODE', 'draft').lower()

# 🌟 NEW LOGIC: Only generate the Reviewer Login and AI Toggle if it is NOT publish mode
admin_controls = ""
if run_mode != 'publish':
    admin_controls = """
    <div class="auth-panel" id="reviewerAccessBox">
        <div class="auth-panel-title">Reviewer Access</div>
        <div class="google-btn-row">
            <div id="googleSignInBtn"></div>
        </div>
        <div id="signedInUser" style="display:none;"></div>
        <div id="authMessage" class="auth-message"></div>
    </div>
    <div class="ai-toggle-wrap" style="display:none;">
        <button type="button" class="ai-toggle-btn" onclick="toggleAIRelevance()">
            Show AI Relevance
        </button>
    </div>
    """

all_articles_html = make_articles_html(all_news, is_finalized=is_finalized)
ist_offset = timezone(timedelta(hours=5, minutes=30))
now_ist = datetime.now(ist_offset)
digit_map = str.maketrans("0123456789", "०१२३४५६७८९")
eng_date, eng_time, hindi_date = now_ist.strftime("%d-%m-%Y"), now_ist.strftime("%H:%M"), now_ist.strftime("%d-%m-%Y").translate(digit_map)
timestamp = f"दिनांक: {hindi_date} • Date: {eng_date} • Time: {eng_time} IST • {len(all_news)} Updates"
build_version = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

# Your full interactive HTML Template (Notice the {admin_controls} variable inside)
html_body = f"""<!DOCTYPE html>
<html data-theme="dark">
<head>
<script src="https://accounts.google.com/gsi/client" async defer></script>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<meta http-equiv="Cache-Control" content="no-cache, no-store, must-revalidate">
<meta http-equiv="Pragma" content="no-cache">
<meta http-equiv="Expires" content="0">
<meta name="build-version" content="{build_version}">
<link rel="stylesheet" href="review-styles.css?v={build_version}">
</head>
<body>
<button class="theme-toggle" id="themeToggle" title="Toggle Theme">☀️</button>
<div class="scroll-container">
    <div class="page-header">
        <img src="./assets/iirs.png" alt="IIRS Logo" class="top-logo left-logo">
        <div class="header-title-wrap">
            # <h2>🌌 अंतरिक्ष समाचार | Space News</h2>
            <h2>🌌 अंतरिक्ष समाचार <span class="title-separator">|</span> <span class="title-eng">Space News</span></h2>
        </div>
        <img src="./assets/ISRO-Color.svg" alt="ISRO Logo" class="top-logo right-logo">
    </div>
    <p style="text-align:center; color:var(--text-secondary); margin-top:12px; margin-bottom:40px;">
        {timestamp}
    </p>
    
    {admin_controls}
    
    {all_articles_html}
    <footer class="footer">
        <div class="footer-text">
            <div>पुस्तकालय एवं सूचना संसाधन प्रभाग द्वारा संकलित, भा.सु.सं.सं</div>
            <div>Compiled by Library and Information Resource Division, IIRS</div>
        </div>
    </footer>
</div>
<div id="customToast" class="custom-toast">
    <div class="toast-title" id="toastTitle"></div>
    <div class="toast-message" id="toastMessage"></div>
</div>
<script src="review-actions.js?v={build_version}"></script>
</body>
</html>
"""

# The safe screen shown to regular employees during Draft Mode
holding_screen_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Space News Digest - Coming Shortly</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; display: flex; justify-content: center; align-items: center; min-height: 100vh; margin: 0; background-color: #f8fafc; color: #334155; text-align: center; }}
        .card {{ background: white; padding: 2.5rem; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); max-width: 480px; border-top: 4px solid #3b82f6; }}
        h1 {{ font-size: 1.5rem; margin-bottom: 0.5rem; color: #0f172a; }}
        p {{ color: #64748b; font-size: 1rem; line-height: 1.6; margin-top: 15px; }}
        .highlight-msg {{ font-weight: bold; color: #1e293b; margin-top: 20px; }}
    </style>
</head>
<body>
    <div class="card">
        <h1>📡 Today's news is coming shortly!</h1>
        <p>Please check back later!</p>
        <p class="highlight-msg">Thanks for cooperation, WE WANT THIS MESSAGE TO DISPLAY</p>
    </div>
</body>
</html>
"""

html_filename = f'Space_News_{datetime.now().strftime("%Y%m%d")}.html'

# File generation conditional based on the run mode
if run_mode == 'publish':
    # PUBLISH: Save real news to main file for the public
    with open(html_filename, 'w', encoding='utf-8') as f:
        f.write(html_body)
    print(f"✅ [PUBLISH] SAVED official public news: {html_filename} with {len(all_news)} items")
else:
    # DRAFT: Hide real news from public, put it in admin-review.html
    with open(html_filename, 'w', encoding='utf-8') as f:
        f.write(holding_screen_html)
    print(f"🔒 [DRAFT] SAVED safe holding screen for the public to {html_filename}")
    
    with open("admin-review.html", "w", encoding="utf-8") as f:
        f.write(html_body)
    print(f"🛠️ [DRAFT] SAVED review tools and full news to admin-review.html")


# =========================
# DOCX Output
# =========================
hindi_days = {"Monday": "सोमवार", "Tuesday": "मंगलवार", "Wednesday": "बुधवार", "Thursday": "गुरुवार", "Friday": "शुक्रवार", "Saturday": "शनिवार", "Sunday": "रविवार"}
eng_day, date_part = now_ist.strftime('%A'), now_ist.strftime('%d/%m/%Y')
hindi_day, hindi_date_part = hindi_days.get(eng_day, ""), date_part.translate(digit_map)
digest_date_str = f"{hindi_day}, {hindi_date_part} | {eng_day}, {date_part}"
docx_filename = f"Space_News_{datetime.now(ist_offset).strftime('%d_%m_%Y')}.docx"

generate_docx(news_items=all_news, output_path=docx_filename, digest_date_str=digest_date_str)
print("📱 HTML + DOCX generation complete.")
