import argparse
import sys
import re
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from zipfile import ZipFile
from copy import deepcopy

from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


# ==========================================
# CONFIGURATION & THEMES
# ==========================================
THEME = {
    "navy": "1F4E78",
    "light_blue": "D9EAF7",
    "dark_text": "1F1F1F",
}

# GH Actions working directory is usually the repo root.
BASE_DIR = Path(".")
TEMP_DIR = BASE_DIR / "temp_build"
OUTPUT_DIR = BASE_DIR


# ==========================================
# STYLING & FORMATTING HELPERS
# ==========================================
def set_run_color(run, hex_color):
    run.font.color.rgb = RGBColor.from_string(hex_color)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")

def set_run_font(run, font_name, size=None, bold=None, italic=None, color=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        set_run_color(run, color)

    r = run._element
    rPr = r.get_or_add_rPr()

    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

def set_cell_margins(cell, top=40, start=40, bottom=40, end=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))


# ==========================================
# TOC & BOOKMARKS
# ==========================================
def make_bookmark_name(num):
    return f"NewsItem_{num}"

def add_bookmark_to_paragraph(paragraph, bookmark_name, bookmark_id):
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, start)
    p.append(end)

def add_internal_hyperlink(paragraph, text, bookmark_name):
    r_begin = OxmlElement("w:r")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fld_char_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' HYPERLINK \\l "{bookmark_name}" '
    r_instr.append(instr)

    r_sep = OxmlElement("w:r")
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fld_char_sep)

    r_text = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r_text.append(rPr)
    
    t = OxmlElement("w:t")
    t.text = text
    r_text.append(t)

    r_end = OxmlElement("w:r")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_char_end)

    for elem in (r_begin, r_instr, r_sep, r_text, r_end):
        paragraph._p.append(elem)

def extract_titles_and_images_from_docx(docx_files):
    items = []
    counter = 1
    toc_img_dir = TEMP_DIR / "extracted_images_for_toc"
    
    for file_path in docx_files:
        doc = Document(str(file_path))
        doc_images = extract_images_from_docx(file_path, toc_img_dir)
        thumb = doc_images[0] if doc_images else None

        for para in doc.paragraphs:
            txt = para.text.strip()
            m = re.match(r"^\d+\.\s+(.+)", txt)
            if m:
                title = m.group(1).strip()
                items.append({
                    "num": counter,
                    "title": title,
                    "image": thumb,
                    "source": file_path.name
                })
                counter += 1
    return items

def add_toc_page(doc, toc_items):
    doc.add_page_break()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(10)
    run = title_para.add_run("Table of Contents")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    total_width = Inches(6.8)
    sno_width = Inches(0.45)
    title_width = total_width - sno_width

    hdr = table.rows[0].cells
    hdr[0].width = sno_width
    hdr[1].width = title_width
    hdr[0].text = "S. No."
    hdr[1].text = "Title"

    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

    for item in toc_items:
        row = table.add_row().cells
        row[0].width = sno_width
        row[1].width = title_width
        row[0].text = str(item["num"])

        p = row[1].paragraphs[0]
        add_internal_hyperlink(p, item["title"], make_bookmark_name(item["num"]))

        for para in row[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows:
        row.cells[0].width = sno_width
        row.cells[1].width = title_width
        for cell in row.cells:
            set_cell_margins(cell, top=35, start=35, bottom=35, end=35)

    doc.add_page_break()

def rebuild_toc_links_and_bookmarks(doc_path):
    doc = Document(str(doc_path))
    final_items = []
    bookmark_id = 1000

    for para in doc.paragraphs:
        txt = para.text.strip()
        m = re.match(r"^(\d+)\.\s+(.+)", txt)
        if m:
            num = int(m.group(1))
            title = m.group(2).strip()
            final_items.append((num, title, para))

    if not final_items:
        doc.save(str(doc_path))
        return

    for num, title, para in final_items:
        add_bookmark_to_paragraph(para, make_bookmark_name(num), bookmark_id)
        bookmark_id += 1

    toc_table = None
    for table in doc.tables:
        try:
            if len(table.rows) >= 1:
                hdr_cells = table.rows[0].cells
                hdr_texts = [c.text.strip().lower() for c in hdr_cells]
                if len(hdr_texts) >= 2 and hdr_texts[0] == "s. no." and hdr_texts[1] == "title":
                    toc_table = table
                    break
        except Exception:
            continue

    if toc_table is None:
        doc.save(str(doc_path))
        return

    while len(toc_table.rows) > 1:
        toc_table._tbl.remove(toc_table.rows[-1]._tr)

    sno_width = Inches(0.45)
    title_width = Inches(6.35)

    for num, title, para in final_items:
        cells = toc_table.add_row().cells
        if len(cells) < 2:
            continue
        cells[0].width = sno_width
        cells[1].width = title_width
        cells[0].text = str(num)
        for p0 in cells[0].paragraphs:
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = cells[1].paragraphs[0]
        p.clear()
        add_internal_hyperlink(p, title, make_bookmark_name(num))
        set_cell_margins(cells[0], top=35, start=35, bottom=35, end=35)
        set_cell_margins(cells[1], top=35, start=35, bottom=35, end=35)

    doc.save(str(doc_path))


# ==========================================
# FILE FETCHING & DATES
# ==========================================
def parse_date_from_filename(name):
    stem = Path(name).stem
    m = re.search(r"Space_News_(\d{2})_(\d{2})_(\d{4})", stem, re.IGNORECASE)
    if not m:
        return None
    return datetime.strptime("_".join(m.groups()), "%d_%m_%Y")

def get_docx_files(folder, from_date, to_date):
    files = [
        p for p in folder.glob("*.docx")
        if not p.name.startswith("~$")
        and not p.name.startswith("SANKALAN_")
    ]
    valid_files = []
    date_pattern = re.compile(r"Space_News_(\d{2}_\d{2}_\d{4})\.docx", re.IGNORECASE)

    for f in files:
        match = date_pattern.search(f.name)
        if match:
            date_str = match.group(1)
            try:
                file_date = datetime.strptime(date_str, "%d_%m_%Y")
                if from_date <= file_date <= to_date:
                    valid_files.append((file_date, f))
            except ValueError:
                continue 

    valid_files.sort(key=lambda x: x[0], reverse=False)
    return [f for _, f in valid_files]


# ==========================================
# IMAGE EXTRACTION & COMPRESSION
# ==========================================
def optimize_docx_images(docx_path):
    """Compresses and downscales images inside the .docx archive to keep file size well under GitHub's 100MB limit."""
    print("⚡ Optimizing and compressing document images...")
    temp_zip_dir = TEMP_DIR / "docx_zip_decompressed"
    if temp_zip_dir.exists():
        shutil.rmtree(temp_zip_dir)
    temp_zip_dir.mkdir(parents=True, exist_ok=True)

    with ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_zip_dir)

    media_dir = temp_zip_dir / "word" / "media"
    if media_dir.exists():
        for img_file in media_dir.iterdir():
            if img_file.suffix.lower() in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp']:
                try:
                    with Image.open(img_file) as im:
                        if im.mode in ("RGBA", "P"):
                            im = im.convert("RGB")
                        
                        max_width = 1200
                        if im.width > max_width:
                            ratio = max_width / im.width
                            new_height = int(im.height * ratio)
                            im = im.resize((max_width, new_height), Image.Resampling.LANCZOS)
                        
                        im.save(img_file, "JPEG", quality=75, optimize=True)
                except Exception as e:
                    print(f"Skipping compression for {img_file.name}: {e}")

    backup_path = docx_path.with_name(docx_path.name + ".bak")
    shutil.move(docx_path, backup_path)
    
    with ZipFile(docx_path, 'w') as zip_out:
        for item in temp_zip_dir.rglob('*'):
            if item.is_file():
                zip_out.write(item, item.relative_to(temp_zip_dir))
    
    backup_path.unlink()
    print("⚡ Image optimization complete!")

def extract_images_from_docx(docx_path, temp_dir):
    images = []
    out_dir = temp_dir / docx_path.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        with ZipFile(docx_path, "r") as z:
            media_files = [n for n in z.namelist() if n.startswith("word/media/")]
            for i, media_name in enumerate(media_files):
                ext = Path(media_name).suffix.lower()
                if ext not in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff", ".webp"]:
                    continue
                out_path = out_dir / f"{i:02d}{ext}"
                with z.open(media_name) as src, open(out_path, "wb") as dst:
                    shutil.copyfileobj(src, dst)
                try:
                    with Image.open(out_path) as im:
                        w, h = im.size
                        if w >= 200 and h >= 150:
                            images.append((w * h, out_path))
                except Exception:
                    pass
    except Exception as e:
        print(f"Could not extract images from {docx_path.name}: {e}")

    images.sort(key=lambda x: x[0], reverse=True)
    return [p for _, p in images]

def collect_top_nine_images(docx_files):
    extract_dir = TEMP_DIR / "extracted_images"
    extract_dir.mkdir(parents=True, exist_ok=True)
    selected = []

    for docx_file in docx_files:
        imgs = extract_images_from_docx(docx_file, extract_dir)
        if imgs:
            selected.append(imgs[0])
        if len(selected) >= 9:
            break

    if len(selected) < 9:
        for docx_file in docx_files:
            imgs = extract_images_from_docx(docx_file, extract_dir)
            for img in imgs[1:]:
                if img not in selected:
                    selected.append(img)
                if len(selected) >= 9:
                    break
            if len(selected) >= 9:
                break
    return selected[:9]

def make_uniform_cover_image(src_path, dst_path, target_size=(900, 650)):
    with Image.open(src_path) as im:
        if im.mode not in ("RGB", "RGBA"):
            im = im.convert("RGB")
        elif im.mode == "RGBA":
            bg = Image.new("RGB", im.size, "white")
            bg.paste(im, mask=im.split()[-1])
            im = bg
        fitted = ImageOps.fit(im, target_size, method=Image.Resampling.LANCZOS, centering=(0.5, 0.5))
        fitted.save(dst_path, format="JPEG", quality=92)

def prepare_uniform_cover_images(image_paths):
    out_dir = TEMP_DIR / "cover_grid_images"
    out_dir.mkdir(parents=True, exist_ok=True)
    prepared = []
    for i, img in enumerate(image_paths[:9], start=1):
        out_path = out_dir / f"cover_{i:02d}.jpg"
        try:
            make_uniform_cover_image(img, out_path)
            prepared.append(out_path)
        except Exception as e:
            print(f"Could not normalize image {img}: {e}")
    return prepared

def format_date_with_ordinal(dt, zero_pad_day=False):
    day = dt.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    day_str = f"{day:02d}" if zero_pad_day else str(day)
    return f"{day_str}{suffix} {dt.strftime('%B %Y')}"

def add_cover_page(doc, image_paths, date_range_text="Fortnightly Compilation"):
    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    top_spacer = doc.add_paragraph()
    top_spacer.paragraph_format.space_before = Pt(0)
    top_spacer.paragraph_format.space_after = Pt(6)

    cover_box = doc.add_table(rows=1, cols=1)
    cover_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_box.autofit = False
    remove_table_borders(cover_box)

    cell = cover_box.cell(0, 0)
    cell.width = Inches(7.1)
    set_cell_shading(cell, "D9EAF7")
    set_cell_margins(cell, top=58, start=80, bottom=55, end=80)

    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("संकलन")
    r0.bold = True
    set_run_font(r0, "Mangal", size=19, bold=True, color=THEME["navy"])

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SANKALAN")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(28)
    set_run_color(r, THEME["navy"])

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(date_range_text)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    set_run_color(r2, THEME["dark_text"])

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run("ई-न्यूज़लेटर | e-Newsletter")
    r3.italic = True
    set_run_font(r3, "Mangal", size=10.5, italic=True, color=THEME["navy"])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(18)

    prepared_images = prepare_uniform_cover_images(image_paths)

    table = doc.add_table(rows=3, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_width = Inches(2.05)
    img_width = Inches(1.85)
    row_height = Inches(1.50)

    idx = 0
    for row in table.rows:
        row.height = row_height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = col_width
            set_cell_margins(cell, top=30, start=30, bottom=30, end=30)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if idx < len(prepared_images):
                try:
                    run = para.add_run()
                    run.add_picture(str(prepared_images[idx]), width=img_width)
                except Exception as e:
                    err_run = para.add_run("[Image error]")
                    set_run_color(err_run, THEME["dark_text"])
            idx += 1

    bottom_spacer = doc.add_paragraph()
    bottom_spacer.paragraph_format.space_before = Pt(14)

    bottom_box = doc.add_table(rows=1, cols=1)
    bottom_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    bottom_box.autofit = False
    remove_table_borders(bottom_box)

    bcell = bottom_box.cell(0, 0)
    bcell.width = Inches(7.1)
    set_cell_shading(bcell, "D9EAF7")
    set_cell_margins(bcell, top=28, start=50, bottom=28, end=50)

    p4 = bcell.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_after = Pt(2)
    r4 = p4.add_run("Indian Institute of Remote Sensing (ISRO), Dehradun")
    r4.bold = True
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(13.5)
    set_run_color(r4, THEME["navy"])

    p5 = bcell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("Collection of recent developments in the realms of Space Technology")
    r5.italic = True
    r5.font.name = "Times New Roman"
    r5.font.size = Pt(12.5)
    set_run_color(r5, THEME["dark_text"])

def create_cover_doc(image_paths, docx_files):
    doc = Document()
    dates = [parse_date_from_filename(f.name) for f in docx_files if parse_date_from_filename(f.name)]
    if dates:
        date_range_text = f"{format_date_with_ordinal(min(dates), zero_pad_day=True)} to {format_date_with_ordinal(max(dates), zero_pad_day=True)}"
    else:
        date_range_text = "Weekly Compilation"
    add_cover_page(doc, image_paths, date_range_text)
    return doc


# ==========================================
# CLEANUP & FORMATTING
# ==========================================
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def delete_table(table):
    tbl = table._element
    tbl.getparent().remove(tbl)
    table._tbl = table._element = None

def is_probable_news_start(block):
    if hasattr(block, "text"):
        txt = block.text.strip()
        if re.match(r"^\d+\.", txt) or txt.lower().startswith("1."):
            return True
    return False

def clean_front_matter_for_later_docx(src_path, out_path):
    doc = Document(str(src_path))
    blocks = list(iter_block_items(doc))
    to_delete = []
    for block in blocks:
        if is_probable_news_start(block):
            break
        to_delete.append(block)
    for block in reversed(to_delete):
        if block.__class__.__name__ == "Paragraph":
            delete_paragraph(block)
        else:
            delete_table(block)
    doc.save(str(out_path))

def renumber_first_level_headings(doc_path):
    doc = Document(str(doc_path))
    counter = 1
    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.match(r"^\d+\.\s", txt):
            new_txt = re.sub(r"^\d+\.\s*", f"{counter}. ", txt, count=1)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_txt
            else:
                para.add_run(new_txt)
            counter += 1
    doc.save(str(doc_path))

def keep_titles_with_following_content(doc_path):
    doc = Document(str(doc_path))
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if re.match(r"^\d+\.\s", txt):
            para.paragraph_format.keep_with_next = True
            para.paragraph_format.keep_together = True
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].paragraph_format.keep_with_next = True
            if i + 2 < len(doc.paragraphs):
                doc.paragraphs[i + 2].paragraph_format.keep_with_next = True
    doc.save(str(doc_path))

def paragraph_has_drawing(para):
    return bool(para._p.xpath(".//w:drawing"))

def add_continuous_section_break(para, num_cols):
    pPr = para._p.get_or_add_pPr()
    sectPr = pPr.first_child_found_in("w:sectPr")
    if sectPr is not None:
        pPr.remove(sectPr)
    sectPr = OxmlElement('w:sectPr')
    type_elm = OxmlElement('w:type')
    type_elm.set(qn('w:val'), 'continuous')
    sectPr.append(type_elm)
    cols_elm = OxmlElement('w:cols')
    cols_elm.set(qn('w:num'), str(num_cols))
    cols_elm.set(qn('w:space'), '720')
    sectPr.append(cols_elm)
    pPr.append(sectPr)

def apply_dynamic_newspaper_columns(doc_path):
    doc = Document(str(doc_path))
    items = []
    current_item = {"title": None, "image": None, "body": []}
    in_news_section = False

    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.match(r"^(\d+)\s*[\.\)]\s+", txt):
            if in_news_section:
                items.append(current_item)
            in_news_section = True
            current_item = {"title": para, "image": None, "body": []}
            continue

        if in_news_section:
            if current_item["image"] is None and paragraph_has_drawing(para):
                current_item["image"] = para
            else:
                current_item["body"].append(para)

    if in_news_section:
        items.append(current_item)

    for idx, item in enumerate(items):
        break_para_1col = item["image"] if item["image"] else item["title"]
        if item["body"]:
            add_continuous_section_break(break_para_1col, num_cols=1)
            if idx < len(items) - 1:
                add_continuous_section_break(item["body"][-1], num_cols=2)

    doc.save(str(doc_path))

def fix_galaxy_emoji_font(doc_path):
    doc = Document(str(doc_path))
    for para in doc.paragraphs:
        if '🌌' not in para.text:
            continue
        original_runs = list(para.runs)
        for r in para._p.xpath('./w:r'):
            para._p.remove(r)
        for run in original_runs:
            if '🌌' in run.text:
                parts = run.text.split('🌌')
                for i, part in enumerate(parts):
                    if part:
                        new_run = para.add_run(part)
                        if run._r.rPr is not None:
                            new_run._r.insert(0, deepcopy(run._r.rPr))
                    if i < len(parts) - 1:
                        emoji_run = para.add_run('🌌')
                        if run._r.rPr is not None:
                            emoji_run._r.insert(0, deepcopy(run._r.rPr))
                        rPr = emoji_run._r.get_or_add_rPr()
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = OxmlElement("w:rFonts")
                            rPr.append(rFonts)
                        for attr in ["ascii", "hAnsi", "cs", "eastAsia"]:
                            rFonts.set(qn(f"w:{attr}"), "Segoe UI Emoji")
            else:
                para._p.append(run._r)
    doc.save(str(doc_path))

def remove_trailing_empty_paragraphs(doc):
    body = doc._element.body
    children = list(body.iterchildren())
    while children:
        last = children[-1]
        if last.tag == qn("w:sectPr") or last.tag != qn("w:p"):
            break
        paras = doc.paragraphs
        if not paras or paras[-1].text.strip() or paras[-1]._p.xpath("./w:pPr/w:sectPr"):
            break
        paras[-1]._element.getparent().remove(paras[-1]._element)
        children = list(body.iterchildren())

def add_disclaimer_page(doc_path):
    doc = Document(str(doc_path))
    remove_trailing_empty_paragraphs(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), '1')
    
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(82)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=115, start=120, bottom=115, end=120)
    set_cell_shading(cell, "D9EAF7")
    set_cell_border(
        cell,
        top={"val": "single", "sz": 10, "color": "9CC2E5", "space": "0"},
        bottom={"val": "single", "sz": 10, "color": "9CC2E5", "space": "0"},
        left={"val": "single", "sz": 10, "color": "9CC2E5", "space": "0"},
        right={"val": "single", "sz": 10, "color": "9CC2E5", "space": "0"},
    )

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1a = p1.add_run("द्वारा अभिकल्पित एवं संकलित\n")
    set_run_font(r1a, "Mangal", size=13, bold=True, color=THEME["navy"])
    r1b = p1.add_run("DESIGNED & COMPILED BY:")
    set_run_font(r1b, "Times New Roman", size=14, bold=True, color=THEME["navy"])

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2a = p2.add_run("पुस्तकालय एवं सूचना संसाधन प्रभाग\n")
    set_run_font(r2a, "Mangal", size=11.5, bold=True, color=THEME["navy"])
    r2b = p2.add_run("Library & Information Resources Division, IIRS (ISRO), Dehradun")
    set_run_font(r2b, "Times New Roman", size=12, color=THEME["dark_text"])

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(40)
    r4a = p4.add_run("अस्वीकरण | ")
    set_run_font(r4a, "Mangal", size=15, bold=True, italic=True, color=THEME["navy"])
    r4b = p4.add_run("Disclaimer:")
    set_run_font(r4b, "Times New Roman", size=16, italic=True, color=THEME["navy"])

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5a = p5.add_run("सभी लेख / सूचनाएँ सार्वजनिक डोमेन में उपलब्ध हैं तथा केवल सूचना एवं सामान्य जागरूकता के उद्देश्य से\nइंटरनेट से संकलित की गई हैं。\n")
    set_run_font(r5a, "Mangal", size=12, bold=True, italic=True, color=THEME["dark_text"])
    r5b = p5.add_run("All the articles / information are available in public domain and are taken from\nInternet for information & general awareness only.")
    set_run_font(r5b, "Times New Roman", size=13, italic=True, color=THEME["dark_text"])

    doc.save(str(doc_path))


# ==========================================
# MASTER MERGE PIPELINE
# ==========================================
def merge_with_cover(docx_files, output_file):
    if not docx_files:
        raise ValueError("No DOCX files found for the given dates.")

    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    cleaned_dir = TEMP_DIR / "cleaned_docx"
    cleaned_dir.mkdir(exist_ok=True)
    temp_body_file = TEMP_DIR / "_temp_body_only.docx"
    temp_final_body_file = TEMP_DIR / "_temp_final_body.docx"

    processed_files = []

    for i, file_path in enumerate(docx_files):
        if i == 0:
            processed_files.append(file_path)
        else:
            cleaned_path = cleaned_dir / file_path.name
            clean_front_matter_for_later_docx(file_path, cleaned_path)
            processed_files.append(cleaned_path)

    print("Creating temporary merged body...")
    master = Document(str(processed_files[0]))
    composer = Composer(master)
    for file_path in processed_files[1:]:
        print(f"Appending: {file_path.name}")
        subdoc = Document(str(file_path))
        composer.append(subdoc)
    composer.save(str(temp_body_file))

    fix_galaxy_emoji_font(str(temp_body_file))
    renumber_first_level_headings(str(temp_body_file))
    rebuild_toc_links_and_bookmarks(str(temp_body_file))
    keep_titles_with_following_content(str(temp_body_file))
    apply_dynamic_newspaper_columns(str(temp_body_file))
    shutil.copy(str(temp_body_file), str(temp_final_body_file))

    filtered_body_docx_files = [temp_final_body_file]
    image_paths = collect_top_nine_images(filtered_body_docx_files)
    toc_items = extract_titles_and_images_from_docx(filtered_body_docx_files)
    cover_doc = create_cover_doc(image_paths, docx_files)
    add_toc_page(cover_doc, toc_items)

    final_composer = Composer(cover_doc)
    filtered_body_doc = Document(str(temp_final_body_file))
    final_composer.append(filtered_body_doc)
    final_composer.save(str(output_file))
    add_disclaimer_page(str(output_file))

    # Compress and squeeze images to stay safely below GitHub's 100MB limit
    optimize_docx_images(output_file)

    print(f"\n✅ Saved merged file to: {output_file}")


# ==========================================
# COMMAND LINE EXECUTION
# ==========================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Weekly Space News Compilation")
    parser.add_argument('--from-date', type=str, default="", help="Start date in YYYY-MM-DD format")
    parser.add_argument('--to-date', type=str, default="", help="End date in YYYY-MM-DD format")
    
    args = parser.parse_args()

    # Intelligent fallback handler for empty/missing dates (7-day rolling window)
    if not args.to_date or args.to_date.strip() == "":
        args.to_date = datetime.now().strftime("%Y-%m-%d")

    if not args.from_date or args.from_date.strip() == "":
        try:
            to_dt_obj = datetime.strptime(args.to_date, "%Y-%m-%d")
        except ValueError:
            to_dt_obj = datetime.now()
        from_dt_obj = to_dt_obj - timedelta(days=7)
        args.from_date = from_dt_obj.strftime("%Y-%m-%d")

    try:
        start_dt = datetime.strptime(args.from_date, "%Y-%m-%d")
        end_dt = datetime.strptime(args.to_date, "%Y-%m-%d")
        
        # Auto-swap if dates are entered in reverse order
        if end_dt < start_dt:
            print("ℹ️ Note: 'To' date is earlier than 'From' date. Automatically swapping them...")
            start_dt, end_dt = end_dt, start_dt
            args.from_date, args.to_date = args.to_date, args.from_date
    except ValueError:
        print("❌ Error: Dates must be in YYYY-MM-DD format.")
        sys.exit(1)

    print(f"📅 Scanning for files between {args.from_date} and {args.to_date}...")
    
    docx_files = get_docx_files(BASE_DIR, start_dt, end_dt)
    
    if not docx_files:
        print(f"❌ No Space_News_*.docx files found in that date range.")
        sys.exit(1)

    # 100% Predictable filename using the exact active date range
    OUTPUT_FILE = OUTPUT_DIR / f"SANKALAN_{args.from_date}_to_{args.to_date}.docx"

    print("📄 Files selected for compilation:")
    for f in docx_files:
        dt = parse_date_from_filename(f.name)
        print(f" - {f.name} | {dt.strftime('%d-%m-%Y') if dt else 'No date parsed'}")

    merge_with_cover(docx_files, output_file=OUTPUT_FILE)
